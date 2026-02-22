"""
Word 文档导出工具 + 物理归档到「我的报告/YYYY-MM-DD/」。

专业排版（纯 python-docx，无需外部模板）：
  - 标题    ：居中、加粗、20pt、深蓝色
  - 二/三级  ：14pt / 12pt 加粗深蓝，段前间距
  - 正文    ：11pt，行距 1.4
  - 免责声明 ：深橙加粗斜体 + 极浅黄底色 + 橙色左边框
  - 页眉    ：右对齐灰色「量化副驾 · 模块名 | 日期」
  - 页脚    ：居中灰色「免责说明 | 第 X 页」
  - 表格    ：深蓝表头白字 + 隔行浅蓝灰底 + Table Grid 框线
"""
from __future__ import annotations

import io
import json
import os
import platform
import re
import subprocess
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── 归档路径 ──────────────────────────────────────────────────────
_PROJECT_ROOT = Path(__file__).parent.parent.parent
ARCHIVE_BASE = _PROJECT_ROOT / "我的报告"

# ── 排版常量 ──────────────────────────────────────────────────────
_F_BODY = 11
_F_H1 = 20
_F_H2 = 14
_F_H3 = 12
_LINE = 1.4

_C_H1 = RGBColor(0x1A, 0x3A, 0x5C)       # 深海蓝
_C_H2 = RGBColor(0x2E, 0x40, 0x57)       # 蓝灰
_C_DISCLM = RGBColor(0xBF, 0x40, 0x00)   # 深橙（免责）
_C_GRAY = RGBColor(0x70, 0x70, 0x70)     # 页眉/页脚灰

_FILL_TH = "2E4057"   # 表头底色（深蓝）
_FILL_TR = "EBF0F5"   # 偶数行底色（极浅蓝灰）
_FILL_WN = "FFF9E6"   # 免责声明底色（极浅黄）
_BORDER_WN = "BF4000" # 免责声明左边框颜色

# 分隔行匹配（|---|---| 格式）
_SEP_RE = re.compile(r"^\s*\|[-:\s|]+\|\s*$")

# Windows 文件名非法字符
_ILLEGAL_WIN_RE = re.compile(r'[\\/:*?"<>|]')


# ══════════════════════════════════════════════════════════════════
# 数据安全工具
# ══════════════════════════════════════════════════════════════════

def safe_text(x) -> str:
    """将任意值安全转换为可写入 docx 的字符串，清除非法控制字符。"""
    if x is None:
        return ""
    if isinstance(x, bool):
        return str(x)
    if isinstance(x, (int, float)):
        return str(x)
    if isinstance(x, (list, dict)):
        try:
            return json.dumps(x, ensure_ascii=False)
        except Exception:
            return str(x)
    s = str(x)
    # 去除 Word 不接受的控制字符（保留 \n \t \r）
    s = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", s)
    return s


def sanitize_filename(name: str) -> str:
    """移除 Windows 非法文件名字符，保留中文；返回安全文件名。"""
    name = _ILLEGAL_WIN_RE.sub("_", name).strip(". ")
    return name or "未命名报告"


@dataclass
class ExportResult:
    """导出操作的统一结果对象。"""
    success: bool
    file_bytes: bytes
    local_path: Path | None
    error_message: str


# ══════════════════════════════════════════════════════════════════
# XML 低级工具
# ══════════════════════════════════════════════════════════════════

def _set_spacing(para, line: float = _LINE, before: float = 0, after: float = 4) -> None:
    """设置段落行距（倍数）和段前段后（pt）。"""
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:line"),     str(int(line * 240)))
    sp.set(qn("w:lineRule"), "auto")
    sp.set(qn("w:before"),   str(int(before * 20)))
    sp.set(qn("w:after"),    str(int(after * 20)))
    pPr.append(sp)


def _shade_para(para, fill: str) -> None:
    """给段落添加背景色，fill 为 6 位十六进制字符串（不含 #）。"""
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:shd")):
        pPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    pPr.append(shd)


def _shade_cell(cell, fill: str) -> None:
    """给表格单元格添加背景色。"""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    tcPr.append(shd)


def _add_para_left_border(para, color_hex: str = _BORDER_WN) -> None:
    """给段落添加左边框（警告/提示块效果）。"""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "12")      # 1.5pt 线宽
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pPr.append(pBdr)


def _add_page_field(para, font_size: int = 8, color: RGBColor = _C_GRAY) -> None:
    """在段落末尾插入 Word PAGE 域（自动页码）。"""
    def _run_with(fldChar_type: str, instr: str | None = None):
        r = para.add_run()
        r.font.size = Pt(font_size)
        r.font.color.rgb = color
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), fldChar_type)
        r._r.append(fc)
        if instr is not None:
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = instr
            r._r.append(it)

    _run_with("begin", " PAGE ")
    _run_with("separate")
    _run_with("end")


# ══════════════════════════════════════════════════════════════════
# 段落级别的内联渲染（处理 **bold** 标记）
# ══════════════════════════════════════════════════════════════════

def _inline(
    para,
    text: str,
    size: float = _F_BODY,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
) -> None:
    """将含 **bold** 标记的文本拆分为多个 run，写入 para。"""
    text = safe_text(text)
    parts = re.split(r"\*\*(.*?)\*\*", text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        # 去掉残留的 *斜体* 包裹（保留文字）
        part = re.sub(r"(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)", r"\1", part)
        run = para.add_run(part)
        run.bold   = bold or bool(idx % 2)
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color


# ══════════════════════════════════════════════════════════════════
# 各类段落添加函数
# ══════════════════════════════════════════════════════════════════

def _add_h1(doc: Document, text: str) -> None:
    text = safe_text(text).strip() or "未命名报告"
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(para, line=1.2, before=6, after=8)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(_F_H1)
    run.font.color.rgb = _C_H1


def _add_h2(doc: Document, text: str) -> None:
    text = safe_text(text)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(para, line=1.2, before=10, after=3)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(_F_H2)
    run.font.color.rgb = _C_H2


def _add_h3(doc: Document, text: str) -> None:
    text = safe_text(text)
    para = doc.add_paragraph()
    _set_spacing(para, line=1.2, before=6, after=2)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(_F_H3)


def _add_body(doc: Document, text: str) -> None:
    text = safe_text(text)
    if not text.strip():
        return
    para = doc.add_paragraph()
    _set_spacing(para, line=_LINE, before=0, after=3)
    _inline(para, text, size=_F_BODY)


def _add_bullet(doc: Document, text: str) -> None:
    text = safe_text(text)
    try:
        para = doc.add_paragraph(style="List Bullet")
    except Exception:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.5)
        text = "· " + text
    _set_spacing(para, line=_LINE, before=0, after=2)
    _inline(para, text, size=_F_BODY)


def _add_numbered(doc: Document, line: str) -> None:
    """有序列表项：直接保留序号，缩进展示。"""
    line = safe_text(line)
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.3)
    _set_spacing(para, line=_LINE, before=0, after=3)
    _inline(para, line, size=_F_BODY)


def _add_rule(doc: Document) -> None:
    para = doc.add_paragraph()
    _set_spacing(para, line=1.0, before=2, after=2)
    run = para.add_run("─" * 40)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def _add_disclaimer(doc: Document, text: str) -> None:
    """免责声明：极浅黄底色 + 橙色左边框 + 深橙加粗斜体。"""
    text = safe_text(re.sub(r"^\*+|\*+$", "", text).strip())
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.3)
    _set_spacing(para, line=1.3, before=6, after=6)
    _shade_para(para, _FILL_WN)
    _add_para_left_border(para, _BORDER_WN)
    run = para.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = _C_DISCLM


def _add_italic_note(doc: Document, text: str) -> None:
    """斜体注释行（灰色小字，比免责声明更低调）。"""
    text = safe_text(re.sub(r"^\*+|\*+$", "", text).strip())
    para = doc.add_paragraph()
    _set_spacing(para, line=_LINE, before=2, after=2)
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _add_quote(doc: Document, line: str) -> None:
    """引用块（> 开头），缩进灰色斜体。"""
    text = safe_text(re.sub(r"^>\s*", "", line).strip())
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.8)
    _set_spacing(para, line=_LINE, before=2, after=2)
    _inline(para, text, size=10, italic=True, color=RGBColor(0x60, 0x60, 0x60))


# ══════════════════════════════════════════════════════════════════
# 表格解析与渲染
# ══════════════════════════════════════════════════════════════════

def _parse_md_table(table_lines: list[str]) -> list[list[str]]:
    """Markdown 表格行 → 二维列表（跳过分隔行）。"""
    rows: list[list[str]] = []
    for line in table_lines:
        if _SEP_RE.match(line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def _add_table(doc: Document, table_lines: list[str]) -> None:
    rows = _parse_md_table(table_lines)
    if not rows:
        _add_body(doc, "（暂无表格数据）")
        return

    n_cols = max(len(r) for r in rows)
    if n_cols == 0:
        _add_body(doc, "（暂无表格数据）")
        return

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        tr = table.rows[r_idx]
        is_header = (r_idx == 0)
        is_even   = (r_idx % 2 == 0)   # 0 = header, 2/4/... = even data rows

        for c_idx in range(n_cols):
            cell = tr.cells[c_idx]
            # safe_text 保证非字符串类型、None 均安全转换
            raw = row_data[c_idx] if c_idx < len(row_data) else ""
            text = safe_text(raw)

            para = cell.paragraphs[0]
            para.clear()
            _set_spacing(para, line=1.2, before=1, after=1)

            if is_header:
                _inline(para, text, size=10, bold=True,
                        color=RGBColor(0xFF, 0xFF, 0xFF))
                _shade_cell(cell, _FILL_TH)
            elif is_even:
                _inline(para, text, size=10)
                _shade_cell(cell, _FILL_TR)
            else:
                _inline(para, text, size=10)

    # 表格后留空行
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════
# 页眉 / 页脚
# ══════════════════════════════════════════════════════════════════

def _set_header_footer(doc: Document, module_label: str) -> None:
    section = doc.sections[0]
    today = date.today().strftime("%Y-%m-%d")

    # ── 页眉（右对齐灰色）────────────────────────────────────────
    header = section.header
    hpara = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hpara.clear()
    hpara.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_spacing(hpara, line=1.0, before=0, after=0)
    run = hpara.add_run(f"量化副驾 · {module_label}  |  {today}")
    run.font.size = Pt(8)
    run.font.color.rgb = _C_GRAY

    # ── 页脚（居中：免责 + 页码）─────────────────────────────────
    footer = section.footer
    fpara = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fpara.clear()
    fpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(fpara, line=1.0, before=0, after=0)

    run1 = fpara.add_run("仅供参考，不构成投资建议  |  第 ")
    run1.font.size = Pt(8)
    run1.font.color.rgb = _C_GRAY

    _add_page_field(fpara)

    run2 = fpara.add_run(" 页")
    run2.font.size = Pt(8)
    run2.font.color.rgb = _C_GRAY


# ══════════════════════════════════════════════════════════════════
# Markdown → docx 主解析器
# ══════════════════════════════════════════════════════════════════

def _markdown_to_docx(doc: Document, content: str) -> None:
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        raw  = lines[i]
        line = raw.rstrip()

        # ── 表格块：连续 | 开头行 ─────────────────────────────
        if line.lstrip().startswith("|"):
            block: list[str] = []
            while i < len(lines) and lines[i].rstrip().lstrip().startswith("|"):
                block.append(lines[i].rstrip())
                i += 1
            _add_table(doc, block)
            continue

        i += 1

        if not line:
            continue

        if line.startswith("# "):
            _add_h1(doc, line[2:].strip())
        elif line.startswith("## "):
            _add_h2(doc, line[3:].strip())
        elif line.startswith("### "):
            _add_h3(doc, line[4:].strip())
        elif line.startswith(("- ", "* ")):
            _add_bullet(doc, line[2:])
        elif re.match(r"^\d+\.\s", line):
            _add_numbered(doc, line)
        elif line.startswith("> "):
            _add_quote(doc, line)
        elif line.startswith("---"):
            _add_rule(doc)
        elif "⚠️" in line:
            _add_disclaimer(doc, line)
        elif (
            line.startswith("*")
            and line.endswith("*")
            and not line.startswith("**")
            and len(line) > 2
        ):
            # *斜体注释* 或 *（注：...）*
            _add_italic_note(doc, line)
        else:
            _add_body(doc, line)


# ══════════════════════════════════════════════════════════════════
# 文档构建
# ══════════════════════════════════════════════════════════════════

def _build_doc(content: str, title: str, module_label: str) -> Document:
    title = safe_text(title).strip() or "未命名报告"
    doc = Document()
    doc.core_properties.author = "light-quant-copilot"

    # 页边距
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    _set_header_footer(doc, module_label)

    # 文档大标题
    _add_h1(doc, title)
    # 标题下日期小字
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(sub, line=1.0, before=0, after=10)
    r = sub.add_run(date.today().strftime("%Y 年 %m 月 %d 日"))
    r.font.size = Pt(10)
    r.font.color.rgb = _C_GRAY

    _markdown_to_docx(doc, content)
    return doc


def _build_doc_bytes(content: str, title: str, module_label: str) -> bytes:
    """构建 Document 并返回 bytes（内存操作，不写磁盘）。"""
    doc = _build_doc(content, title, module_label)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════
# 归档目录
# ══════════════════════════════════════════════════════════════════

def _today_dir() -> Path:
    d = ARCHIVE_BASE / date.today().strftime("%Y-%m-%d")
    d.mkdir(parents=True, exist_ok=True)
    return d


# ══════════════════════════════════════════════════════════════════
# 统一导出入口
# ══════════════════════════════════════════════════════════════════

def export_report_to_docx(
    content: str,
    title: str,
    filename: str,
    module_label: str,
) -> ExportResult:
    """
    统一导出函数：生成 bytes + 落盘归档，返回 ExportResult。

    - 数据自动过 safe_text 清洗
    - 文件名去除 Windows 非法字符，保留中文
    - 同名文件冲突时追加时间戳另存
    - 捕获所有异常，返回含错误信息的 ExportResult，不崩页面
    """
    try:
        # 数据清洗
        content = safe_text(content) or "（内容为空）"
        title = safe_text(title).strip() or "未命名报告"

        # 生成字节（内存）
        file_bytes = _build_doc_bytes(content, title, module_label)

        # 落盘归档
        safe_fn = sanitize_filename(filename)
        if not safe_fn.endswith(".docx"):
            safe_fn += ".docx"
        target = _today_dir() / safe_fn

        # 同名冲突：追加时间戳
        if target.exists():
            ts = datetime.now().strftime("%H%M%S")
            target = target.parent / f"{target.stem}_{ts}{target.suffix}"

        try:
            target.write_bytes(file_bytes)
        except PermissionError:
            # 文件被占用：再追加时间戳另存一份
            ts = datetime.now().strftime("%H%M%S")
            target = target.parent / f"{target.stem}_{ts}{target.suffix}"
            target.write_bytes(file_bytes)

        return ExportResult(
            success=True,
            file_bytes=file_bytes,
            local_path=target,
            error_message="",
        )

    except PermissionError as exc:
        return ExportResult(
            success=False, file_bytes=b"", local_path=None,
            error_message=f"写入失败：文件被其他程序占用或无磁盘写权限（{exc}）",
        )
    except OSError as exc:
        return ExportResult(
            success=False, file_bytes=b"", local_path=None,
            error_message=f"文件系统错误：路径不可写或磁盘已满（{exc}）",
        )
    except Exception as exc:
        return ExportResult(
            success=False, file_bytes=b"", local_path=None,
            error_message=f"导出失败（内容异常或 python-docx 错误）：{exc}",
        )


# ══════════════════════════════════════════════════════════════════
# 三模块快捷导出（调用 export_report_to_docx）
# ══════════════════════════════════════════════════════════════════

def export_morning_report(content: str) -> ExportResult:
    today = date.today().strftime("%Y-%m-%d")
    return export_report_to_docx(
        content,
        title="今日晨报",
        filename=f"{today}_晨报.docx",
        module_label="晨报",
    )


def export_risk_report(content: str) -> ExportResult:
    today = date.today().strftime("%Y-%m-%d")
    return export_report_to_docx(
        content,
        title="风控清单",
        filename=f"{today}_风控清单.docx",
        module_label="风控",
    )


def export_review(content: str) -> ExportResult:
    today = date.today().strftime("%Y-%m-%d")
    return export_report_to_docx(
        content,
        title="收盘复盘",
        filename=f"{today}_收盘复盘.docx",
        module_label="复盘",
    )


def export_test_doc() -> ExportResult:
    """生成测试用 docx，用于验证导出链路（python-docx、下载按钮、本地落盘）。"""
    content = """\
## 测试段落

这是一段**加粗**文字，用于验证 python-docx 基本渲染。

- 列表项一：中文内容正常
- 列表项二：数字 123 / 小数 3.14
- 列表项三：English & symbols !@#

| 股票代码 | 名称 | 当前价格 | 涨跌幅 |
|---------|------|---------|-------|
| 600519 | 贵州茅台 | 1800.00 | +1.2% |
| 000001 | 平安银行 | 12.34 | -0.5% |
| 300750 | 宁德时代 | 230.00 | +3.0% |

> 引用块：仅供测试，不构成任何投资建议。

### 三级标题

1. 有序列表第一项
2. 有序列表第二项

⚠️ 免责说明：本文档由系统自动生成，仅用于测试 Word 导出功能，请勿用于实际投资决策。
"""
    return export_report_to_docx(
        content,
        title="导出测试文档",
        filename="test_export.docx",
        module_label="测试",
    )


# ══════════════════════════════════════════════════════════════════
# 文件夹打开工具（供 Streamlit UI 调用）
# ══════════════════════════════════════════════════════════════════

def open_folder(folder: Path) -> str | None:
    """
    用系统默认文件管理器打开文件夹。
    返回 None 表示成功；返回文件夹路径字符串表示失败（供 UI 显示）。
    """
    try:
        if platform.system() == "Windows":
            os.startfile(str(folder))
        else:
            result = subprocess.run(
                ["xdg-open", str(folder)],
                timeout=3,
                capture_output=True,
            )
            if result.returncode != 0:
                return str(folder)
        return None
    except Exception:
        return str(folder)
